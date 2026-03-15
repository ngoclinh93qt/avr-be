"""Export Research Brief as PDF or Word document."""

import io
import logging
from datetime import datetime

from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import StreamingResponse

from app.core.supabase_client import supabase_service
from app.api.deps import get_current_user_id

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/export", tags=["export"])


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _safe(value, fallback: str = "—") -> str:
    if value is None:
        return fallback
    return str(value).strip() or fallback


def _format_date(iso_str: str) -> str:
    try:
        return datetime.fromisoformat(iso_str.replace("Z", "+00:00")).strftime("%d/%m/%Y")
    except Exception:
        return iso_str[:10] if iso_str else "—"


# ─── PDF generator ────────────────────────────────────────────────────────────

def _build_pdf(session: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    )
    from reportlab.lib.enums import TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    PRIMARY = HexColor("#2563EB")
    MUTED = HexColor("#6B7280")

    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=16, textColor=PRIMARY, spaceAfter=4)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12, textColor=PRIMARY, spaceBefore=14, spaceAfter=4)
    body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=9, leading=14, spaceAfter=4)
    small = ParagraphStyle("Small", parent=styles["Normal"], fontSize=8, textColor=MUTED, leading=12)
    label = ParagraphStyle("Label", parent=styles["Normal"], fontSize=8, textColor=MUTED, spaceAfter=1)

    story = []

    # ── Header ──
    story.append(Paragraph("✦ AVR — Research Brief", h1))
    story.append(Paragraph(
        f"Tạo: {_format_date(session.get('created_at', ''))} · "
        f"Session: {session.get('id', '')[:8]}",
        small
    ))
    story.append(HRFlowable(width="100%", thickness=1, color=PRIMARY, spaceAfter=10))

    # ── 1. Blueprint ──
    blueprint = session.get("blueprint") or {}
    story.append(Paragraph("1. Research Blueprint", h2))
    bp_data = [
        ["Câu hỏi NC", _safe(blueprint.get("intervention_or_exposure"))],
        ["Thiết kế", _safe(blueprint.get("design_type", "")).replace("_", " ").title()],
        ["Đối tượng", f"{_safe(blueprint.get('population'))} (n = {_safe(blueprint.get('sample_size'))})"],
        ["Kết cục chính", _safe(blueprint.get("primary_outcome"))],
        ["Setting", _safe(blueprint.get("setting"))],
    ]
    if blueprint.get("statistical_approach"):
        bp_data.append(["Phân tích", _safe(blueprint.get("statistical_approach"))])

    tbl = Table(bp_data, colWidths=[4 * cm, 12 * cm])
    tbl.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("TEXTCOLOR", (0, 0), (0, -1), MUTED),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [HexColor("#F9FAFB"), None]),
        ("GRID", (0, 0), (-1, -1), 0.25, HexColor("#E5E7EB")),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(tbl)

    # ── 2. Novelty Check ──
    novelty = session.get("novelty_check") or {}
    if novelty:
        story.append(Paragraph("2. Kiểm tra độ mới (PubMed)", h2))
        story.append(Paragraph(
            f"~{novelty.get('count', 0)} bài tương tự · "
            f"Keywords: {', '.join(novelty.get('keywords_used', []))}",
            label
        ))
        for i, p in enumerate(novelty.get("papers", [])[:5], 1):
            story.append(Paragraph(
                f"{i}. <b>{_safe(p.get('title'))}</b> — "
                f"{_safe(p.get('authors'))} ({_safe(p.get('year'))}) — {_safe(p.get('journal'))}",
                body
            ))
        if novelty.get("commentary"):
            story.append(Spacer(1, 4))
            story.append(Paragraph(f"💡 {novelty['commentary']}", body))

    # ── 3. Estimated Abstract ──
    abstract = session.get("estimated_abstract")
    if abstract:
        story.append(Paragraph("3. Abstract ước tính", h2))
        for line in abstract.split("\n"):
            if line.strip():
                story.append(Paragraph(line.strip(), body))

    # ── 4. Journal Suggestions ──
    journals = session.get("journal_suggestions") or []
    if journals:
        story.append(Paragraph("4. Gợi ý tạp chí", h2))
        for i, j in enumerate(journals[:3], 1):
            name = _safe(j.get("name"))
            if_ = f"IF {j['impact_factor']:.1f}" if j.get("impact_factor") else ""
            oa = _safe(j.get("open_access"), "")
            story.append(Paragraph(
                f"{i}. <b>{name}</b> · {if_} · {oa}",
                body
            ))

    # ── 5. Research Roadmap ──
    roadmap = session.get("roadmap") or {}
    if roadmap:
        story.append(Paragraph("5. Lộ trình nghiên cứu", h2))
        story.append(Paragraph(
            f"Checklist: {_safe(roadmap.get('checklist_type'))} · "
            f"Tổng thời gian: {_safe(roadmap.get('total_timeline_estimate'))}",
            label
        ))
        for step in roadmap.get("steps", []):
            story.append(Paragraph(
                f"<b>Bước {step['step_number']}: {_safe(step.get('title'))}</b> "
                f"({_safe(step.get('duration_estimate'))}) — {_safe(step.get('who'))}",
                body
            ))
            story.append(Paragraph(_safe(step.get("description")), small))

    # ── Footer ──
    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", thickness=0.5, color=MUTED))
    story.append(Paragraph(
        "Tài liệu này do AVR Research Mentor tạo tự động. "
        "Luôn kiểm tra yêu cầu của tạp chí trước khi submit.",
        small
    ))

    doc.build(story)
    return buf.getvalue()


# ─── Word generator ───────────────────────────────────────────────────────────

def _build_docx(session: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    BLUE = RGBColor(0x25, 0x63, 0xEB)
    GRAY = RGBColor(0x6B, 0x72, 0x80)

    def add_heading(text: str, level: int = 1):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = BLUE

    def add_para(text: str, bold: bool = False, color: RGBColor | None = None, size: int = 9):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        return p

    def add_kv(label: str, value: str):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.font.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = GRAY
        r2 = p.add_run(value)
        r2.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(2)

    # ── Header ──
    title = doc.add_paragraph()
    run = title.add_run("✦ AVR — Research Brief")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BLUE

    add_para(
        f"Tạo: {_format_date(session.get('created_at', ''))} · Session: {session.get('id', '')[:8]}",
        color=GRAY, size=8
    )

    # ── 1. Blueprint ──
    add_heading("1. Research Blueprint", level=1)
    blueprint = session.get("blueprint") or {}
    add_kv("Câu hỏi NC", _safe(blueprint.get("intervention_or_exposure")))
    add_kv("Thiết kế", _safe(blueprint.get("design_type", "")).replace("_", " ").title())
    add_kv("Đối tượng", f"{_safe(blueprint.get('population'))} (n = {_safe(blueprint.get('sample_size'))})")
    add_kv("Kết cục chính", _safe(blueprint.get("primary_outcome")))
    add_kv("Setting", _safe(blueprint.get("setting")))

    # ── 2. Novelty Check ──
    novelty = session.get("novelty_check") or {}
    if novelty:
        add_heading("2. Kiểm tra độ mới (PubMed)", level=1)
        add_para(
            f"~{novelty.get('count', 0)} bài tương tự · Keywords: {', '.join(novelty.get('keywords_used', []))}",
            color=GRAY, size=8
        )
        for i, p in enumerate(novelty.get("papers", [])[:5], 1):
            add_para(
                f"{i}. {_safe(p.get('title'))} — {_safe(p.get('authors'))} ({_safe(p.get('year'))}) — {_safe(p.get('journal'))}",
                size=8
            )
        if novelty.get("commentary"):
            add_para(f"💡 {novelty['commentary']}", size=9)

    # ── 3. Estimated Abstract ──
    abstract = session.get("estimated_abstract")
    if abstract:
        add_heading("3. Abstract ước tính", level=1)
        for line in abstract.split("\n"):
            if line.strip():
                add_para(line.strip())

    # ── 4. Journal Suggestions ──
    journals = session.get("journal_suggestions") or []
    if journals:
        add_heading("4. Gợi ý tạp chí", level=1)
        for i, j in enumerate(journals[:3], 1):
            name = _safe(j.get("name"))
            if_ = f"IF {j['impact_factor']:.1f}" if j.get("impact_factor") else ""
            oa = _safe(j.get("open_access"), "")
            add_para(f"{i}. {name} · {if_} · {oa}")

    # ── 5. Research Roadmap ──
    roadmap = session.get("roadmap") or {}
    if roadmap:
        add_heading("5. Lộ trình nghiên cứu", level=1)
        add_para(
            f"Checklist: {_safe(roadmap.get('checklist_type'))} · "
            f"Tổng thời gian: {_safe(roadmap.get('total_timeline_estimate'))}",
            color=GRAY, size=8
        )
        for step in roadmap.get("steps", []):
            add_para(
                f"Bước {step['step_number']}: {_safe(step.get('title'))} "
                f"({_safe(step.get('duration_estimate'))}) — {_safe(step.get('who'))}",
                bold=True
            )
            add_para(_safe(step.get("description")), color=GRAY, size=8)

    # ── Footer ──
    doc.add_paragraph()
    footer = add_para(
        "Tài liệu này do AVR Research Mentor tạo tự động. "
        "Luôn kiểm tra yêu cầu của tạp chí trước khi submit.",
        color=GRAY, size=8
    )
    footer.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Endpoint ─────────────────────────────────────────────────────────────────

@router.get("/brief")
async def export_brief(
    session_id: str = Query(...),
    format: str = Query("pdf", pattern="^(pdf|docx)$"),
    user_id: str = Depends(get_current_user_id),
):
    """
    Export Research Brief as PDF or Word document.

    GET /api/v1/export/brief?session_id=...&format=pdf
    GET /api/v1/export/brief?session_id=...&format=docx
    """
    session = await supabase_service.get_research_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if session.get("user_id") != user_id:
        raise HTTPException(status_code=403, detail="Access denied")

    if not session.get("estimated_abstract"):
        raise HTTPException(
            status_code=400,
            detail="Abstract not yet generated. Complete Phase 1 first."
        )

    try:
        if format == "pdf":
            content = _build_pdf(session)
            media_type = "application/pdf"
            filename = f"AVR_Research_Brief_{session_id[:8]}.pdf"
        else:
            content = _build_docx(session)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"AVR_Research_Brief_{session_id[:8]}.docx"

    except ImportError as e:
        raise HTTPException(
            status_code=501,
            detail=f"Export library not installed: {e}. Run: pip install reportlab python-docx"
        )
    except Exception as e:
        logger.exception("Export failed for session %s", session_id)
        raise HTTPException(status_code=500, detail=f"Export failed: {e}")

    return StreamingResponse(
        io.BytesIO(content),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
